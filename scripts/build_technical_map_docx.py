"""Generate TECHNICAL_MAP.docx (EN + မြန်မာ) for the NovaLink MM Mini App."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def cell_shading(cell, hex_rgb: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_rgb)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_path = root / "TECHNICAL_MAP.docx"

    doc = Document()

    title = doc.add_heading("NovaLink MM · Q-Practice Style Technical Map", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "Telegram Mini App — React + TypeScript + Vite + Firebase (Firestore + Storage). "
        "Document generated from repo structure.",
        italic=True,
    )
    doc.add_paragraph()

    # Legend
    add_heading(doc, "Legend — မြင်သာအောင် အရောင်ခွဲ", 1)
    legend_rows = [
        ("2563EB", "User UI", "အသုံးပြုသူ Mini App (browser SPA)"),
        ("7C3AED", "Admin UI", "Firebase Console · approve · key ထည့် (repo ထဲ admin web မပါ)"),
        ("EA580C", "Payment", "ငွေပေးချေ · screenshot · verify flow"),
        ("64748B", "Legacy / External", "Optional NovaLink Python bot / သီးခြား server"),
        ("16A34A", "Firestore data", "Cloud orders · counters · vpnKeys"),
        ("CA8A04", "Local / cache", "Browser localStorage (u5_* keys)"),
        ("0891B2", "React FE files", ".tsx · hooks · components (Dart မဟုတ်)"),
        ("475569", "Config / JSON / Rules", "firebase.json · rules · indexes"),
    ]
    table = doc.add_table(rows=1 + len(legend_rows), cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "■"
    hdr[1].text = "Area (EN)"
    hdr[2].text = "မြန်မာ"
    for i, (hex_c, en, mm) in enumerate(legend_rows, start=1):
        row = table.rows[i].cells
        row[0].text = "■"
        cell_shading(row[0], hex_c)
        if row[0].paragraphs[0].runs:
            row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        row[1].text = en
        row[2].text = mm

    doc.add_paragraph()

    # At-a-glance
    add_heading(doc, "At-a-glance summary", 1)
    t2 = doc.add_table(rows=6, cols=3)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Topic"
    t2.rows[0].cells[1].text = "EN"
    t2.rows[0].cells[2].text = "မြန်မာ"
    rows2 = [
        ("Frontend Apps", "1 SPA — purchase, keys, FAQ, guides (no in-repo admin app).", "User-facing တစ်ခုတည်း။"),
        ("Primary DB", "Firestore: orders, meta/counters, vpnKeys (optional).", "အော်ဒါနှင့် keys လမ်းကြောင်း။"),
        ("Object storage", "Firebase Storage: payments/{userId}/… screenshots.", "ငွေလွှဲ screenshot ဖိုင်များ။"),
        ("Hosting", "Firebase Hosting or Vultr VPS + nginx (static dist/).", "Firebase သို့ VPS မှာ dist တင်။"),
        ("Backup", "Manual export ready; auto Scheduler needs GCP Billing.", "လက်ဖြင့် export၊ auto က billing လို။"),
    ]
    for i, (a, b, c) in enumerate(rows2, start=1):
        t2.rows[i].cells[0].text = a
        t2.rows[i].cells[1].text = b
        t2.rows[i].cells[2].text = c

    doc.add_paragraph()

    # System map text
    add_heading(doc, "System map — text flow", 1)
    flow = """User (Telegram) → HTTPS Hosting → React SPA (frontend/src)
  → data/store/appStore.ts (local first + sync)
  → backend/firebase/* → Firestore (orders, meta) + Storage (payments/)
Admin: Firebase Console edits same Firestore documents.
Optional: external Python bot if same Firebase project."""
    doc.add_paragraph(flow)

    # Area table
    add_heading(doc, "Area — what / source of truth / key files", 1)
    t3 = doc.add_table(rows=7, cols=4)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "Area"
    t3.rows[0].cells[1].text = "What (EN)"
    t3.rows[0].cells[2].text = "Source of truth"
    t3.rows[0].cells[3].text = "Key files"
    areas = [
        ("User UI", "Regions, packages, payment, verify, orders, keys, FAQ.", "Firestore + localStorage", "frontend/src/pages/*, App.tsx"),
        ("Admin UI", "Manual approve; set status + accessUrl.", "Firebase Console", "— (Console)"),
        ("Payment", "Reference + screenshot upload.", "orders + Storage", "PaymentPage, PaymentVerifyPage, backend/firebase/storage.ts"),
        ("Firestore", "Orders, counters, vpnKeys.", "Cloud", "backend/firebase/orders.ts, keys.ts"),
        ("Local cache", "Draft + orders mirror.", "localStorage", "data/store/appStore.ts"),
        ("Rules/JSON", "Security + hosting + indexes.", "Repo → deploy", "firebase/*.rules, firestore.indexes.json, firebase.json"),
    ]
    for i, row in enumerate(areas, start=1):
        for j, val in enumerate(row):
            t3.rows[i].cells[j].text = val

    doc.add_paragraph()

    # Functions
    add_heading(doc, "Important functions", 1)
    t4 = doc.add_table(rows=7, cols=4)
    t4.style = "Table Grid"
    headers = ["Function", "Purpose (EN)", "Purpose (MM)", "Where"]
    for j, h in enumerate(headers):
        t4.rows[0].cells[j].text = h
    funcs = [
        ("saveOrder", "Local-first, then Firestore.", "local သိမ်းပြီး sync။", "data/store/appStore.ts"),
        ("createOrderId", "Transaction counter or local fallback.", "counter သို့ fallback。", "appStore.ts / orders.ts"),
        ("uploadPaymentScreenshot", "Data URL → Storage URL.", "screenshot upload。", "storage.ts"),
        ("getOrders / getOrder", "Remote or local read.", "အော်ဒါဖတ်。", "appStore.ts"),
        ("getActiveKeys", "Completed orders + vpnKeys.", "keys ပေါင်း。", "appStore.ts"),
        ("isFirebaseConfigured", "Gate SDK via env.", "env စစ်。", "backend/firebase/client.ts"),
    ]
    for i, row in enumerate(funcs, start=1):
        for j, val in enumerate(row):
            t4.rows[i].cells[j].text = val

    doc.add_paragraph()

    # Data storage
    add_heading(doc, "Data storage", 1)
    t5 = doc.add_table(rows=6, cols=4)
    t5.style = "Table Grid"
    t5.rows[0].cells[0].text = "Data"
    t5.rows[0].cells[1].text = "Stored in"
    t5.rows[0].cells[2].text = "Path / Collection"
    t5.rows[0].cells[3].text = "Notes"
    data_rows = [
        ("Orders", "Firestore", "orders/{telegramUserId}_{orderId}", "Merged fields from Order type."),
        ("Counters", "Firestore", "meta/counters orderId", "Transaction."),
        ("Screenshots", "Storage", "payments/{userId}/...", "Not stored as blob in Firestore."),
        ("vpnKeys", "Firestore", "vpnKeys/*", "Optional."),
        ("Cache", "localStorage", "u5_orders, u5_purchase_draft, u5_keys", "Offline resilience."),
    ]
    for i, row in enumerate(data_rows, start=1):
        for j, val in enumerate(row):
            t5.rows[i].cells[j].text = val

    doc.add_paragraph()

    # Stack
    add_heading(doc, "Language & stack", 1)
    t6 = doc.add_table(rows=6, cols=4)
    t6.style = "Table Grid"
    t6.rows[0].cells[0].text = "Layer"
    t6.rows[0].cells[1].text = "Language"
    t6.rows[0].cells[2].text = "Framework"
    t6.rows[0].cells[3].text = "Deploy"
    stack_rows = [
        ("Mini App UI", "TypeScript", "React 19 + Vite", "Firebase Hosting / VPS"),
        ("Firebase client", "TypeScript", "Firebase JS SDK", "Bundled"),
        ("Cloud DB", "—", "Firestore", "Firebase project"),
        ("Files", "—", "Cloud Storage", "Same project"),
        ("Rules", "Rules DSL", "firestore.rules, storage.rules", "firebase deploy"),
    ]
    for i, row in enumerate(stack_rows, start=1):
        for j, val in enumerate(row):
            t6.rows[i].cells[j].text = val

    doc.add_paragraph()

    # Key React files
    add_heading(doc, "Key frontend / data / backend files", 1)
    t7 = doc.add_table(rows=13, cols=3)
    t7.style = "Table Grid"
    t7.rows[0].cells[0].text = "Path"
    t7.rows[0].cells[1].text = "Role (EN)"
    t7.rows[0].cells[2].text = "မြန်မာ"
    files = [
        ("frontend/src/App.tsx", "Routes, Telegram back button, bottom menu", "လမ်းကြောင်း၊ Telegram"),
        ("frontend/src/pages/PaymentVerifyPage.tsx", "Reference + screenshot submit", "အတည်ပြု တင်ပြမှု"),
        ("data/store/appStore.ts", "Orders + Firebase bridge", "အော်ဒါ orchestration"),
        ("data/config.ts", "Regions, prices, FAQ", "config အချက်အလက်"),
        ("backend/firebase/orders.ts", "Firestore orders API", "orders ရေးဖတ်"),
        ("backend/firebase/storage.ts", "Screenshot upload", "Storage upload"),
        ("backend/firebase/client.ts", "Firebase init", "SDK စတင်"),
        ("firebase/firestore.rules", "Security rules", "စည်းမျဉ်း"),
        ("firebase/storage.rules", "Storage rules", "Storage စည်းမျဉ်း"),
        ("firebase/firestore.indexes.json", "Composite indexes", "index များ"),
        ("vite.config.ts", "Build + path aliases @, @data, @backend", "build config"),
        ("package.json", "npm scripts (dev, build, deploy)", "scripts"),
    ]
    for i, row in enumerate(files, start=1):
        for j, val in enumerate(row):
            t7.rows[i].cells[j].text = val

    doc.add_paragraph()

    # Backup
    add_heading(doc, "Backup & ops", 1)
    doc.add_paragraph(
        "EN: Use Firestore export (Console or gcloud) to GCS; Storage can be synced with gsutil. "
        "Automated nightly exports require GCP billing and Cloud Scheduler.\n\n"
        "MM: Firestore ကို Console သို့ gcloud မှ export လုပ်ပါ။ နေ့စဉ် auto သည် billing လိုအပ်သည်။"
    )

    # Timeline
    add_heading(doc, "Timeline — ကာလ (conceptual phases)", 1)
    phases = [
        "Phase 1 — Setup: repo, .env VITE_FIREBASE_*, npm run dev.",
        "Phase 2 — Firebase: deploy firestore rules + indexes; enable Storage + deploy storage rules.",
        "Phase 3 — Hosting: npm run build; deploy to Firebase Hosting or copy dist to Vultr nginx.",
        "Phase 4 — BotFather: set Mini App URL (HTTPS).",
        "Phase 5 — Ops: admin uses Console to complete orders + accessUrl; optional backups.",
    ]
    for p in phases:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_paragraph()
    add_para(doc, f"Output file: {out_path.name}", italic=True)

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
